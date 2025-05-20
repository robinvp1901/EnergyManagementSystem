import plotly.graph_objects as go
from plotly.subplots import make_subplots
from plotly.express import colors
# --- Export figuren naar Word-document ---
from docx import Document
from docx.shared import Inches
import os
import pandas as pd

export_map = "./export/"
# Laad CSV
df = pd.read_csv(export_map + "resultaat_SIM_1_Controle_zonder_PV_en_Batterij.csv", parse_dates=['timestamp'])
# Laad tweede DataFrame voor SIM_2
df2 = pd.read_csv(export_map + "resultaat_SIM_2_Controle_alleen_3_PV-panelen.csv", parse_dates=['timestamp'])
# Laad DataFrames voor SIM 3, 4 en 5
df3 = pd.read_csv(export_map + "resultaat_SIM_3_Alleen_2.4_kWh_Batterij.csv", parse_dates=['timestamp'])
df4 = pd.read_csv(export_map + "resultaat_SIM_4_Alleen_3.5_kWh_Batterij.csv", parse_dates=['timestamp'])
df5 = pd.read_csv(export_map + "resultaat_SIM_5_Alleen_4.8_kWh_Batterij.csv", parse_dates=['timestamp'])
 # Laad DataFrames voor SIM 6, 7 en 8
df6 = pd.read_csv(export_map + "resultaat_SIM_6_3_PV-Panelen_+_2.4_kWH_Batterij.csv", parse_dates=['timestamp'])
df7 = pd.read_csv(export_map + "resultaat_SIM_7_3_PV-Panelen_+_3.5_kWH_Batterij.csv", parse_dates=['timestamp'])
df8 = pd.read_csv(export_map + "resultaat_SIM_8_3_PV-Panelen_+_4.8_kWH_Batterij.csv", parse_dates=['timestamp'])
# Laad DataFrames voor SIM 9, 10 en 11
df9 = pd.read_csv(export_map + "resultaat_SIM_9_6_PV-Panelen_+_2.4_kWH_Batterij.csv", parse_dates=['timestamp'])
df10 = pd.read_csv(export_map + "resultaat_SIM_10_6_PV-Panelen_+_3.5_kWH_Batterij.csv", parse_dates=['timestamp'])
df11 = pd.read_csv(export_map + "resultaat_SIM_11_6_PV-Panelen_+_4.8_kWH_Batterij.csv", parse_dates=['timestamp'])


# Samenvattingsstatistieken laden
summary_df = pd.read_csv(export_map + "samenvatting_scenario's.csv")

palette = colors.qualitative.Plotly


def make_line_plot(x, y_data: dict, title, y_title, style_dict=None, yaxis_range=None, yaxis_type='linear', xaxis_range=None):
    """
    Maakt een uniforme lijnplot met standaardinstellingen.

    Parameters:
    - x: x-as waarden (bijv. tijd)
    - y_data: dict met {label: y-waarden}
    - title: grafiektitel
    - x_title, y_title: assenlabels
    """
    fig = go.Figure()

    for i, (label, y) in enumerate(y_data.items()):
        style = style_dict.get(label, {}) if style_dict else {}
        color = style.get("color", palette[i % len(palette)])
        dash = style.get("dash", "solid")
        width = style.get("width", 2)
        fill = 'tozeroy' if label == 'PV Geleverd [kW]' else None
        fillcolor = style.get("fillcolor", None)

        fig.add_trace(go.Scatter(
            x=x, y=y,
            name=label,
            line=dict(color=color, dash=dash, width=width),
            legendgroup=str(i),
            fill=fill,
            fillcolor=fillcolor,
        ))

    fig.update_layout(
        title=title,
        xaxis_title="Tijdstap",
        yaxis_title=y_title,
        yaxis=dict(
            range=yaxis_range,
            type=yaxis_type,
            ticklabelposition="outside", ticklen=8, tickwidth=1, tickcolor='black', ticklabelstandoff=10
        ),
        # Layoutopties voor afstand tussen assen en ticklabels
        xaxis=dict(range=xaxis_range, ticklabelposition="outside", ticklen=8, tickwidth=1, tickcolor='black', ticklabelstandoff=10),
        template='plotly_white',
        width=800,
        height=500,
        font=dict(size=14),
        margin=dict(l=40, r=40, t=40, b=40),
        legend=dict(x=1.02, y=1, bgcolor='rgba(255,255,255,0.5)', orientation='v')
    )

    return fig


# Nieuwe functie voor heatmap-subplots
def make_heatmap_subplots(df, fields, title):
    """
    Maakt subplot met heatmaps per veld (bijv. Load, Grid Power).

    Parameters:
    - df: DataFrame met kolom 'timestamp' en gevraagde energievelden
    - fields: lijst met kolomnamen zoals ['load [kW]', 'grid_power [kW]']
    - title: bovenste titel voor alle subplots

    Retourneert: plotly.graph_objects.Figure
    """
    df['dag'] = df['timestamp'].dt.date
    df['uur'] = df['timestamp'].dt.hour

    n = len(fields)
    custom_titles = {
        'load [kW]': 'Verbruik [kWh]',
        'grid_power [kW]': 'Net inkoop [kWh]',
        'pv_used [kW]': 'PV Geleverd [kWh]',
        'soc_kwh [kWh]': 'Batterij [kWhh]'
    }
    subplot_titles = [custom_titles.get(f, f) for f in fields]

    fig = make_subplots(
        rows=n, cols=1,
        shared_xaxes=True,
        subplot_titles=subplot_titles,
        vertical_spacing=0.10
    )

    for i, veld in enumerate(fields):
        data_matrix = df.pivot(index='uur', columns='dag', values=veld)
        # Sorteer index zodat uren in natuurlijke volgorde staan
        data_matrix = data_matrix.sort_index()
        show_colorbar = True
        # Kies kleurenpalet en schaal per veld
        if veld == 'soc_kwh [kWh]':
            colorscale = 'RdBu_r'
            zmin = 0
            zmax = 8
        else:
            colorscale = 'Bluyl'
            zmin = 0
            zmax = 2.5
        heatmap = go.Heatmap(
            z=data_matrix.values,
            x=data_matrix.columns.astype(str),
            y=[f"{uur:02d}:00" for uur in data_matrix.index],
            colorscale=colorscale,
            showscale=show_colorbar,
            colorbar=dict(
                title=dict(
                    text=subplot_titles[i],
                    side='right'
                ),
                len=0.6,
                y=0,
                yanchor='bottom',
                x=1 + i * 0.2,
                xanchor='left',
                thickness=10,
            ) if show_colorbar else None,
            zsmooth=False,
            zmin=zmin,
            zmax=zmax
        )
        fig.add_trace(heatmap, row=i+1, col=1)
        # Zorg dat elke subplot zijn eigen x-as ticks toont
        fig.update_xaxes(showticklabels=True, row=i+1, col=1)

    fig.update_layout(
        height=300 * n,
        width=800,
        title=title,
        font=dict(size=16),
        template='plotly_white',
        margin=dict(l=20, r=20, t=80, b=40),
        legend=dict(x=1.02, y=1, bgcolor='rgba(255,255,255,0.5)', orientation='v')
    )
    # Zet x-as titel voor onderste subplot
    fig.update_xaxes(title_text='Datum', row=n, col=1)
    # Zet y-as titels voor alle subplots
    for i in range(1, n + 1):
        fig.update_yaxes(title_text='Tijdstip', row=i, col=1, autorange=True)
    # Layoutopties voor afstand tussen assen en ticklabels
    fig.update_xaxes(ticklabelposition="outside", ticklen=8, tickwidth=1, tickcolor='black')
    fig.update_yaxes(ticklabelposition="outside", ticklen=8, tickwidth=1, tickcolor='black')
    fig.update_xaxes(ticklabelstandoff=10)
    fig.update_yaxes(ticklabelstandoff=10)

    for i in range(1, n + 1):
        fig.update_yaxes(title_font=dict(size=14), tickfont=dict(size=12), row=i, col=1)
    for i in range(1, n + 1):
        fig.update_xaxes(title_font=dict(size=14), tickfont=dict(size=12), row=i, col=1)

    symbool_map = {
        'load [kW]': '⚡️',
        'grid_power [kW]':'🔌',
        'pv_used [kW]': '☀️',
        'pv_avail [kW]': '☀️'
    }
    annotations = []
    for field in fields:
        symbool = symbool_map.get(field, '')
        naam = custom_titles.get(field, field)
        # Toon totaal voor alle velden behalve soc_kwh
        if field != 'soc_kwh [kWh]':
            totaal = df[field].sum()
            annotations.append(f"<b>{symbool} {naam}:</b> {totaal:.1f} kWh")
    # Voeg Grid → Load toe
    if 'grid_to_load [kW]' in df.columns:
        grid_to_load = df['grid_to_load [kW]'].sum()
        annotations.append(f"<b>🔌 Grid → Load [kWh]:</b> {grid_to_load:.1f} kWh")
    # Voeg net naar batterij toe
    if 'grid_to_bat [kW]' in df.columns:
        grid_to_bat = df['grid_to_bat [kW]'].sum()
        annotations.append(f"<b>🔌 Net → Batterij [kWh]:</b> {grid_to_bat:.1f} kWh")
    # Voeg netkosten toe indien van toepassing
    if 'grid_power [kW]' in fields and 'price [€/kWh]' in df.columns:
        kosten_net = (df['grid_power [kW]'] * df['price [€/kWh]']).sum()
        annotations.append(f"<b>💰 Inkoopkosten:</b> €{kosten_net:.1f}")
    # Voeg batterijlevering aan load toe
    if 'battery_to_load [kW]' in df.columns:
        battery_to_load = df['battery_to_load [kW]'].sum()
        annotations.append(f"<b>🔋 Batterij → Load [kWh]:</b> {battery_to_load:.1f} kWh")

    # Voeg PV-statistieken toe indien van toepassing
    if 'pv_avail [kW]' in df.columns:
        annotations.append(f"<b>☀️ PV beschikbaar:</b> {df['pv_avail [kW]'].sum():.1f} kWh")
    if 'pv_used [kW]' in df.columns:
        annotations.append(f"<b>☀️ PV geleverd:</b> {df['pv_used [kW]'].sum():.1f} kWh")
    if 'pv_to_load [kW]' in df.columns:
        annotations.append(f"<b>☀️→📦️ PV → Load:</b> {df['pv_to_load [kW]'].sum():.1f} kWh")
    if 'pv_to_bat [kW]' in df.columns:
        annotations.append(f"<b>☀️→🔋 PV → Batterij:</b> {df['pv_to_bat [kW]'].sum():.1f} kWh")

    annotation_text = "<br>".join(annotations)
    fig.add_annotation(
        text=annotation_text,
        showarrow=False,
        xref="paper", yref="paper",
        x=1.4, y=0.9,
        align="left",
        font=dict(size=12),
        bordercolor="black",
        borderwidth=1,
        bgcolor="white",
        opacity=0.9
    )

    return fig


style = {
    'Verbruik [kW]': {'color': palette[1], 'dash': 'solid', 'width': 1.5},
    'Net [kW]': {'color': palette[2], 'dash': 'solid', 'width': 1.5, 'fillcolor': 'rgba(44, 160, 44, 0.5)'},
    'Prijs [€/kWh]': {'color': palette[9], 'dash': 'solid', 'width': 1},
    'PV Beschikbaar [kW]': {'color': 'grey', 'dash': 'dot', 'width': 1.5},
    'PV Geleverd [kW]': {
        'color': 'rgba(255, 127, 14, 1)',
        'fillcolor': 'rgba(255, 127, 14, 0.5)',
        'dash': 'solid',
        'width': 1.5
    },
    # Toevoegen stijlen voor Net SIM 3, 4, 5
    'Net SIM 3 [kW]': {'color': palette[2], 'dash': 'solid', 'width': 1.5, 'fillcolor': 'rgba(44, 160, 44, 0.5)'},
    'Net SIM 4 [kW]': {'color': palette[2], 'dash': 'solid', 'width': 1.5, 'fillcolor': 'rgba(44, 160, 44, 0.5)'},
    'Net SIM 5 [kW]': {'color': palette[2], 'dash': 'solid', 'width': 1.5, 'fillcolor': 'rgba(44, 160, 44, 0.5)'}
}


# Heatmap subplot toevoegen en opslaan
heatmap_fig = make_heatmap_subplots(df, ['load [kW]', 'grid_power [kW]'], 'SIM 1: Verbruik en inkoop net - Heatmaps')

# Heatmap subplot voor SIM 2: Verbruik, Net en PV Geleverd
heatmap_fig_sim2 = make_heatmap_subplots(
    df2,
    ['load [kW]', 'grid_power [kW]', 'pv_used [kW]'],
    'SIM 2: Verbruik, Net en PV Geleverd - Heatmaps'
)

# Heatmap subplots voor SIM 3, 4 en 5
make_heatmap_subplots(
    df3,
    ['load [kW]', 'grid_power [kW]', 'soc_kwh [kWh]'],
    'SIM 3: Verbruik, Net en Batterij - Heatmaps'
)

make_heatmap_subplots(
    df4,
    ['load [kW]', 'grid_power [kW]', 'soc_kwh [kWh]'],
    'SIM 4: Verbruik, Net en Batterij - Heatmaps'
)

make_heatmap_subplots(
    df5,
    ['load [kW]', 'grid_power [kW]', 'soc_kwh [kWh]'],
    'SIM 5: Verbruik, Net en Batterij - Heatmaps'
)

# Heatmap subplots voor SIM 6, 7 en 8
make_heatmap_subplots(
    df6,
    ['load [kW]', 'grid_power [kW]', 'soc_kwh [kWh]'],
    'SIM 6: Verbruik, Net en Batterij - Heatmaps'
)

make_heatmap_subplots(
    df7,
    ['load [kW]', 'grid_power [kW]', 'soc_kwh [kWh]'],
    'SIM 7: Verbruik, Net en Batterij - Heatmaps'
)

make_heatmap_subplots(
    df8,
    ['load [kW]', 'grid_power [kW]', 'soc_kwh [kWh]'],
    'SIM 8: Verbruik, Net en Batterij - Heatmaps'
)

# Heatmap subplots voor SIM 9, 10 en 11 NA de definitie van make_heatmap_subplots
make_heatmap_subplots(
    df9,
    ['load [kW]', 'grid_power [kW]', 'soc_kwh [kWh]'],
    'SIM 9: Verbruik, Net en Batterij - Heatmaps'
)
make_heatmap_subplots(
    df10,
    ['load [kW]', 'grid_power [kW]', 'soc_kwh [kWh]'],
    'SIM 10: Verbruik, Net en Batterij - Heatmaps'
)
make_heatmap_subplots(
    df11,
    ['load [kW]', 'grid_power [kW]', 'soc_kwh [kWh]'],
    'SIM 11: Verbruik, Net en Batterij - Heatmaps'
)

fig1 = go.Figure()

fig1.add_trace(go.Bar(
    x=df['timestamp'],
    y=df['load [kW]'],
    name='Verbruik [kW]',
    marker_color=style['Verbruik [kW]']['color']
))

fig1.update_layout(
    title='SIM 1: Verbruik (3 dagen)',
    xaxis_title="Tijdstap",
    yaxis_title="Vermogen [kWh]",
    yaxis=dict(range=[0, 2.5], type='linear'),
    xaxis=dict(range=[df['timestamp'].min(), df['timestamp'].min() + pd.Timedelta(days=3)]),
    template='plotly_white',
    width=800,
    height=500,
    font=dict(size=14),
    margin=dict(l=40, r=40, t=40, b=40),
    legend=dict(x=1.02, y=1, bgcolor='rgba(255,255,255,0.5)', orientation='v')
)

fig2 = make_line_plot(
    x=df['timestamp'],
    y_data={
        'Prijs [€/kWh]': df['price [€/kWh]']
    },
    title='Energieprijs 2022 incl. belastingen en accijns',
    y_title='Prijs [€/kWh]',
    yaxis_range=[0, 1.5],
    yaxis_type='linear',
    style_dict=style  # <-- deze regel toevoegen
)


fig3 = make_subplots(
    rows=2, cols=1,
    shared_xaxes=True,
    vertical_spacing=0.1,
    subplot_titles=["Verbruik & Net inkoop", "PV Beschikbaar & Geleverd"]
)

# Bovenste subplot: verbruik en net
fig3.add_trace(go.Bar(
    x=df2['timestamp'],
    y=df2['load [kW]'],
    name='Verbruik [kWh]',
    marker_color=style['Verbruik [kW]']['color']
), row=1, col=1)

fig3.add_trace(go.Bar(
    x=df2['timestamp'],
    y=df2['grid_power [kW]'],
    name='Net inkoop [kWh]',
    marker_color=style['Net [kW]']['color']
), row=1, col=1)

# Onderste subplot: pv beschikbaar en geleverd
fig3.add_trace(go.Bar(
    x=df2['timestamp'],
    y=df2['pv_avail [kW]'],
    name='PV Beschikbaar [kWh]',
    marker_color=style['PV Beschikbaar [kW]']['color']
), row=2, col=1)

fig3.add_trace(go.Bar(
    x=df2['timestamp'],
    y=df2['pv_used [kW]'],
    name='PV Geleverd [kWh]',
    marker_color=style['PV Geleverd [kW]']['color']
), row=2, col=1)

fig3.update_layout(
    title='SIM 2: Vraag, Net en 3 PV-panelen',
    height=600,
    width=800,
    xaxis_title="Tijdstap per uur",
    yaxis_title="Vermogen [kWh]",
    template='plotly_white',
    font=dict(size=14),
    margin=dict(l=40, r=40, t=60, b=40),
    legend=dict(x=1.02, y=1, bgcolor='rgba(255,255,255,0.5)', orientation='v'),
    barmode='overlay'
)

fig3.update_yaxes(title_text='Vermogen [kWh]', row=1, col=1)
fig3.update_yaxes(title_text='Vermogen [kWh]', row=2, col=1)
fig3.update_xaxes(range=[df2['timestamp'].min(), df2['timestamp'].min() + pd.Timedelta(days=3)])
# Zet x-as titel onder de tweede subplot
fig3.update_xaxes(title_text="Tijdstap per uur", row=2, col=1)
fig3.update_xaxes(title_text=None, showticklabels=True, row=1, col=1)


fig4 = make_subplots(
    rows=4, cols=1,
    shared_xaxes=True,
    vertical_spacing=0.08,
    subplot_titles=["SIM 3: Verbruik + Net + Batterij 2.4 kWh",
                    "SIM 4: Verbruik + Net + Batterij 3.5 kWh",
                    "SIM 5: Verbruik + Net + Batterij 4.8 kWh",
                    "Prijs [€/kWh]"]
)


fig4.add_trace(go.Scatter(
    x=df3['timestamp'],
    y=df3['soc_kwh [kWh]'],
    name='Batterij',
    line=dict(
        color=palette[0], dash='dot', width=2
    ),
    showlegend=True,
    legendgroup='Batterij SOC',
), row=1, col=1)

fig4.add_trace(go.Bar(
    x=df3['timestamp'],
    y=df3['grid_power [kW]'],
    name='Net inkoop [kWh]',
    marker_color=style['Net SIM 3 [kW]']['color']
), row=1, col=1)

fig4.add_trace(go.Bar(
    x=df3['timestamp'],
    y=df3['load [kW]'],
    name='Verbruik [kWh]',
    marker_color=style['Verbruik [kW]']['color']
), row=1, col=1)

# SIM 4 (row=2): showlegend=False
fig4.add_trace(go.Scatter(
    x=df4['timestamp'],
    y=df4['soc_kwh [kWh]'],
    name='Batterij 3.5 kWh',
    line=dict(
        color=palette[0], dash='dot', width=2
    ),
    showlegend=False,
    legendgroup='Batterij',
), row=2, col=1)

fig4.add_trace(go.Bar(
    x=df4['timestamp'],
    y=df4['grid_power [kW]'],
    name='Net inkoop [kWh]',
    marker_color=style['Net SIM 4 [kW]']['color'],
    showlegend=False,
), row=2, col=1)

fig4.add_trace(go.Bar(
    x=df4['timestamp'],
    y=df4['load [kW]'],
    name='Verbruik [kWh]',
    marker_color=style['Verbruik [kW]']['color'],
    showlegend=False,
), row=2, col=1)

# SIM 5 (row=3): showlegend=False
fig4.add_trace(go.Scatter(
    x=df5['timestamp'],
    y=df5['soc_kwh [kWh]'],
    name='Batterij 4.8 kWh',
    line=dict(
        color=palette[0], dash='dot', width=2
    ),
    showlegend=False,
    legendgroup='Batterij',
), row=3, col=1)

fig4.add_trace(go.Bar(
    x=df5['timestamp'],
    y=df5['grid_power [kW]'],
    name='Net inkoop [kWh]',
    marker_color=style['Net SIM 5 [kW]']['color'],
    showlegend=False,
), row=3, col=1)

fig4.add_trace(go.Bar(
    x=df5['timestamp'],
    y=df5['load [kW]'],
    name='Verbruik [kWh]',
    marker_color=style['Verbruik [kW]']['color'],
    showlegend=False,
), row=3, col=1)

# Prijs [€/kWh] trace onderaan (row=4)
fig4.add_trace(
    go.Scatter(
        x=df3['timestamp'],
        y=df3['price [€/kWh]'],
        name='Prijs [€/kWh]',
        line=dict(
            color=style['Prijs [€/kWh]']['color'],
            dash=style['Prijs [€/kWh]']['dash'],
            width=style['Prijs [€/kWh]']['width']
        ),
        showlegend=False
    ),
    row=4, col=1
)

fig4.update_layout(
    title="SIM 3–5: Vergelijking Verbruik, Net en Batterij",
    height=1200,
    width=800,
    xaxis_title="Tijdstap per uur",
    yaxis_title="Vermogen [kW] / SOC [kWh]",
    template='plotly_white',
    font=dict(size=14),
    margin=dict(l=40, r=40, t=60, b=40),
    legend=dict(x=1.02, y=1, bgcolor='rgba(255,255,255,0.5)', orientation='v'),
    barmode='overlay'
)
# Beperk de x-as tot de eerste 10 dagen
fig4.update_xaxes(range=[df['timestamp'].min(), df['timestamp'].min() + pd.Timedelta(days=10)])
# Voeg individuele y-as titels toe per subplot
for i in range(1, 4):
    fig4.update_yaxes(title_text="Vermogen [kW] / SOC [kWh]", row=i, col=1)
    fig4.update_xaxes(showticklabels=True, row=i, col=1)
# Prijs subplot y-as titel
fig4.update_yaxes(title_text="Prijs [€/kWh]", row=4, col=1)
fig4.update_xaxes(title_text="Tijdstap per uur", row=4, col=1)
fig4.update_xaxes(title_text="", row=1, col=1)


fig5 = make_subplots(
    rows=4, cols=1,
    shared_xaxes=True,
    vertical_spacing=0.08,
    subplot_titles=[
        "SIM 6: Verbruik + Net + PV + Batterij 2.4 kWh",
        "SIM 7: Verbruik + Net + PV + Batterij 3.5 kWh",
        "SIM 8: Verbruik + Net + PV + Batterij 4.8 kWh",
        "Prijs [€/kWh]"
    ]
)

# SIM 6 (row=1)
fig5.add_trace(go.Scatter(
    x=df6['timestamp'],
    y=df6['soc_kwh [kWh]'],
    name='Batterij [kWh]',
    line=dict(color=palette[0], dash='dot', width=2),
    showlegend=True,
    legendgroup='Batterij'
), row=1, col=1)

fig5.add_trace(go.Bar(
    x=df6['timestamp'],
    y=df6['grid_power [kW]'],
    name='Net inkoop [kWh]',
    marker_color=style['Net SIM 3 [kW]']['color'],
    showlegend=True,
), row=1, col=1)

fig5.add_trace(go.Bar(
    x=df6['timestamp'],
    y=df6['load [kW]'],
    name='Verbruik [kWh]',
    marker_color=style['Verbruik [kW]']['color'],
    showlegend=True,
), row=1, col=1)

# SIM 7 (row=2)
fig5.add_trace(go.Scatter(
    x=df7['timestamp'],
    y=df7['soc_kwh [kWh]'],
    name='Batterij 3.5 kWh',
    line=dict(color=palette[0], dash='dot', width=2),
    showlegend=False,
    legendgroup='Batterij'
), row=2, col=1)
fig5.add_trace(go.Bar(
    x=df7['timestamp'],
    y=df7['grid_power [kW]'],
    name='Net SIM 7 [kWh]',
    marker_color=style['Net SIM 3 [kW]']['color'],
    showlegend=False,
), row=2, col=1)
fig5.add_trace(go.Bar(
    x=df7['timestamp'],
    y=df7['load [kW]'],
    name='Verbruik [kWh]',
    marker_color=style['Verbruik [kW]']['color'],
    showlegend=False,
), row=2, col=1)

# SIM 8 (row=3)
fig5.add_trace(go.Scatter(
    x=df8['timestamp'],
    y=df8['soc_kwh [kWh]'],
    name='Batterij 4.8 kWh',
    line=dict(color=palette[0], dash='dot', width=2),
    showlegend=False,
    legendgroup='Batterij'
), row=3, col=1)
fig5.add_trace(go.Bar(
    x=df8['timestamp'],
    y=df8['grid_power [kW]'],
    name='Net SIM 8 [kWh]',
    marker_color=style['Net SIM 3 [kW]']['color'],
    showlegend=False,
), row=3, col=1)
fig5.add_trace(go.Bar(
    x=df8['timestamp'],
    y=df8['load [kW]'],
    name='Verbruik [kWh]',
    marker_color=style['Verbruik [kW]']['color'],
    showlegend=False,
), row=3, col=1)

# Prijs [€/kWh] trace onderaan (row=4)
fig5.add_trace(
    go.Scatter(
        x=df6['timestamp'],
        y=df6['price [€/kWh]'],
        name='Prijs [€/kWh]',
        line=dict(
            color=style['Prijs [€/kWh]']['color'],
            dash=style['Prijs [€/kWh]']['dash'],
            width=style['Prijs [€/kWh]']['width']
        ),
        showlegend=False
    ),
    row=4, col=1
)

fig5.update_layout(
    title="SIM 6–8: Vergelijking Verbruik, Net en Batterij",
    height=1200,
    width=800,
    xaxis_title="Datum",
    yaxis_title="Vermogen [kW] / SOC [kWh]",
    template='plotly_white',
    font=dict(size=14),
    margin=dict(l=40, r=40, t=60, b=40),
    legend=dict(x=1.02, y=1, bgcolor='rgba(255,255,255,0.5)', orientation='v'),
    barmode='overlay'
)
fig5.update_xaxes(range=[df['timestamp'].min(), df['timestamp'].min() + pd.Timedelta(days=10)])
for i in range(1, 4):
    fig5.update_yaxes(title_text="Vermogen [kW] / SOC [kWh]", row=i, col=1)
    fig5.update_xaxes(showticklabels=True, row=i, col=1)
fig5.update_xaxes(title_text="Datum", row=3, col=1)
fig5.update_xaxes(title_text="", row=1, col=1)
fig5.update_yaxes(title_text="Prijs [€/kWh]", row=4, col=1)
fig5.update_xaxes(title_text="Datum", row=4, col=1)



fig6 = make_subplots(
    rows=3, cols=1,
    shared_xaxes=True,
    vertical_spacing=0.1,
    subplot_titles=[
        "SIM 6: PV Beschikbaar en Geleverd",
        "SIM 7: PV Beschikbaar en Geleverd",
        "SIM 8: PV Beschikbaar en Geleverd"
    ]
)

for idx, (dfi, sim) in enumerate([(df6, 6), (df7, 7), (df8, 8)]):
    row = idx + 1
    for label in ['PV Beschikbaar [kW]', 'PV Geleverd [kW]']:
        y_ = dfi['pv_avail [kW]'] if label == 'PV Beschikbaar [kW]' else dfi['pv_used [kW]']
        fig6.add_trace(go.Bar(
            x=dfi['timestamp'],
            y=y_,
            name=label,
            marker_color=style[label]['color'],
            showlegend=(row == 1),
            legendgroup=label
        ), row=row, col=1)

fig6.update_layout(
    title="SIM 6–8: 3 PV-Panelen + Batterij 2.4/3.5/4.8 kWh",
    height=900,
    width=800,
    xaxis_title="Datum",
    yaxis_title="Vermogen [kWh]",
    template='plotly_white',
    font=dict(size=14),
    margin=dict(l=40, r=40, t=60, b=40),
    legend=dict(x=1.02, y=1, bgcolor='rgba(255,255,255,0.5)', orientation='v'),
    barmode='overlay'
)

for i in range(1, 4):
    fig6.update_yaxes(title_text="Vermogen [kWh]", row=i, col=1)
    fig6.update_xaxes(showticklabels=True, row=i, col=1)
fig6.update_xaxes(title_text="Datum", row=3, col=1)
fig6.update_xaxes(title_text="", row=1, col=1)
# Beperk de x-as tot de eerste 10 dagen
fig6.update_xaxes(range=[df6['timestamp'].min(), df6['timestamp'].min() + pd.Timedelta(days=5)])



fig7 = make_subplots(
    rows=4, cols=1,
    shared_xaxes=True,
    vertical_spacing=0.08,
    subplot_titles=[
        "SIM 9: Verbruik + Net + PV + Batterij 2.4 kWh",
        "SIM 10: Verbruik + Net + PV + Batterij 3.5 kWh",
        "SIM 11: Verbruik + Net + PV + Batterij 4.8 kWh",
        "Prijs [€/kWh]"
    ]
)

for idx, (dfi, title, net_label) in enumerate([
    (df9, 'Batterij [kWh]', 'Net inkoop [kW]'),
    (df10, 'Batterij 3.5 kWh', 'Net SIM 6 [kW]'),
    (df11, 'Batterij 4.8 kWh', 'Net SIM 6 [kW]')
]):
    row = idx + 1
    fig7.add_trace(go.Scatter(
        x=dfi['timestamp'],
        y=dfi['soc_kwh [kWh]'],
        name=title,
        line=dict(color=palette[0], dash='dot', width=2),
        showlegend=(row == 1),
        legendgroup='Batterij'
    ), row=row, col=1)
    fig7.add_trace(go.Bar(
        x=dfi['timestamp'],
        y=dfi['grid_power [kW]'],
        name=net_label,
        marker_color=style['Net SIM 3 [kW]']['color'],
        showlegend=(row == 1),
        legendgroup='Net'
    ), row=row, col=1)
    fig7.add_trace(go.Bar(
        x=dfi['timestamp'],
        y=dfi['load [kW]'],
        name='Verbruik [kWh]',
        marker_color=style['Verbruik [kW]']['color'],
        showlegend=(row == 1),
        legendgroup='Verbruik [kW]'
    ), row=row, col=1)

# Prijs [€/kWh] trace onderaan (row=4)
fig7.add_trace(
    go.Scatter(
        x=df9['timestamp'],
        y=df9['price [€/kWh]'],
        name='Prijs [€/kWh]',
        line=dict(
            color=style['Prijs [€/kWh]']['color'],
            dash=style['Prijs [€/kWh]']['dash'],
            width=style['Prijs [€/kWh]']['width']
        ),
        showlegend=False
    ),
    row=4, col=1
)

fig7.update_layout(
    title="SIM 9–11: Vergelijking Verbruik, Net en Batterij",
    height=1200,
    width=800,
    xaxis_title="Datum",
    yaxis_title="Vermogen [kWh] / SOC [kWh]",
    template='plotly_white',
    font=dict(size=14),
    margin=dict(l=40, r=40, t=60, b=40),
    legend=dict(x=1.02, y=1, bgcolor='rgba(255,255,255,0.5)', orientation='v'),
    barmode='overlay'
)
fig7.update_xaxes(range=[df['timestamp'].min(),
                         df['timestamp'].min() + pd.Timedelta(days=5)]
                  )

for i in range(1, 4):
    fig7.update_yaxes(title_text="Vermogen [kWh] / SOC [kWh]", row=i, col=1)
    fig7.update_xaxes(showticklabels=True, row=i, col=1)
fig7.update_xaxes(title_text="Datum", row=3, col=1)
fig7.update_xaxes(title_text="", row=1, col=1)
fig7.update_yaxes(title_text="Prijs [€/kWh]", row=4, col=1)
fig7.update_xaxes(title_text="Datum", row=4, col=1)





fig8 = make_subplots(
    rows=3, cols=1,
    shared_xaxes=True,
    vertical_spacing=0.1,
    subplot_titles=[
        "SIM 9: PV Beschikbaar en Geleverd",
        "SIM 10: PV Beschikbaar en Geleverd",
        "SIM 11: PV Beschikbaar en Geleverd"
    ]
)

for idx, (dfi, sim) in enumerate([(df9, 9), (df10, 10), (df11, 11)]):
    row = idx + 1
    for label in ['PV Beschikbaar [kW]', 'PV Geleverd [kW]']:
        y_ = dfi['pv_avail [kW]'] if label == 'PV Beschikbaar [kW]' else dfi['pv_used [kW]']
        fig8.add_trace(go.Bar(
            x=dfi['timestamp'],
            y=y_,
            name=label,
            marker_color=style[label]['color'],
            showlegend=(row == 1),
            legendgroup=label
        ), row=row, col=1)

fig8.update_layout(
    title="SIM 9–11: 6 PV-Panelen + Batterij 2.4/3.5/4.8 kWh",
    height=900,
    width=800,
    xaxis_title="Datum",
    yaxis_title="Vermogen [kWh]",
    template='plotly_white',
    font=dict(size=14),
    margin=dict(l=40, r=40, t=60, b=40),
    legend=dict(x=1.02, y=1, bgcolor='rgba(255,255,255,0.5)', orientation='v'),
    barmode='overlay'
)

for i in range(1, 3+1):
    fig8.update_yaxes(title_text="Vermogen [kWh]", row=i, col=1)
    fig8.update_xaxes(showticklabels=True, row=i, col=1)
fig8.update_xaxes(title_text="Datum", row=3, col=1)
fig8.update_xaxes(title_text="", row=1, col=1)

# Beperk de x-as tot de eerste 10 dagen
fig8.update_xaxes(range=[df9['timestamp'].min(),
                         df9['timestamp'].min() + pd.Timedelta(days=5)]
                  )




# Maak export directory aan
doc_export_path = "./export/figures_docx/"
os.makedirs(doc_export_path, exist_ok=True)

figures = [
    ("fig1", fig1),
    ("heatmap_fig", heatmap_fig),
    ("fig2", fig2),
    ("fig3", fig3),
    ("heatmap_fig_sim2", heatmap_fig_sim2),
    ("fig4", fig4),
    ("fig5", fig5),
    ("fig6", fig6),
    ("fig7", fig7),
    ("fig8", fig8),

]

# Voeg heatmaps voor SIM 3 t/m 11 toe
for sim_num, df_sim in zip(range(3, 12), [df3, df4, df5, df6, df7, df8, df9, df10, df11]):
    sim_title = f"SIM {sim_num}: Verbruik, Net en Batterij - Heatmaps"
    sim_fig = make_heatmap_subplots(
        df_sim,
        ['load [kW]', 'grid_power [kW]', 'soc_kwh [kWh]'],
        sim_title
    )
    figures.append((f"heatmap_fig_sim{sim_num}", sim_fig))

image_paths = []
for name, fig in figures:
    path = os.path.join(doc_export_path, f"{name}.png")
    fig.write_image(path, width=800, height=1200, scale=1)
    image_paths.append((name, path))

# Maak Word-document aan
doc = Document()
doc.add_heading("Visualisatie Resultaten Energie Management Systeem", level=1)

for name, img_path in image_paths:
    doc.add_paragraph(name.upper(), style='Heading2')
    doc.add_picture(img_path, width=Inches(6))
    doc.add_paragraph("")

# Opslaan als DOCX
doc_path = os.path.join(doc_export_path, "Energie_Visualisatie_Resultaten.docx")
doc.save(doc_path)
print(f"Document opgeslagen als: {doc_path}")